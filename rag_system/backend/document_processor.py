"""
Document processing module for extracting text from PDF, DOCX, and TXT files.
Handles chunking and metadata extraction.
"""

import os
from pathlib import Path
from typing import List, Dict, Tuple
import logging

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats and extract text content."""
    
    CHUNK_SIZE = 500  # Characters per chunk
    OVERLAP = 100  # Character overlap between chunks
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF file."""
        try:
            text = ""
            metadata = {"filename": Path(file_path).name, "format": "pdf", "pages": 0}
            
            with open(file_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                metadata["pages"] = len(reader.pages)
                
                for page in reader.pages:
                    text += page.extract_text()
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX file."""
        try:
            text = ""
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            metadata = {
                "filename": Path(file_path).name,
                "format": "docx",
                "paragraphs": len(doc.paragraphs)
            }
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, Dict]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            
            metadata = {
                "filename": Path(file_path).name,
                "format": "txt",
                "lines": len(text.split('\n'))
            }
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    @classmethod
    def extract_text(cls, file_path: str) -> Tuple[str, Dict]:
        """Extract text from any supported format."""
        file_path = str(file_path)
        suffix = Path(file_path).suffix.lower()
        
        if suffix == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif suffix == '.docx':
            return cls.extract_text_from_docx(file_path)
        elif suffix == '.txt':
            return cls.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, 
                   overlap: int = OVERLAP) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                if last_period > chunk_size // 2:
                    end = start + last_period + 1
            
            chunks.append(text[start:end].strip())
            start = end - overlap
        
        return [c for c in chunks if c]  # Filter empty chunks
    
    @classmethod
    def process_document(cls, file_path: str) -> Dict:
        """Process a document and return chunks with metadata."""
        text, metadata = cls.extract_text(file_path)
        chunks = cls.chunk_text(text)
        
        return {
            "filename": metadata.get("filename", ""),
            "format": metadata.get("format", ""),
            "chunks": chunks,
            "chunk_count": len(chunks),
            "text_length": len(text),
            "metadata": metadata
        }
    
    @classmethod
    def process_documents(cls, directory: str) -> List[Dict]:
        """Process all documents in a directory."""
        documents = []
        
        if not os.path.exists(directory):
            logger.warning(f"Directory {directory} does not exist")
            return documents
        
        for file_path in Path(directory).glob("*"):
            if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                try:
                    doc_data = cls.process_document(str(file_path))
                    documents.append(doc_data)
                    logger.info(f"Processed {file_path.name}: {doc_data['chunk_count']} chunks")
                except Exception as e:
                    logger.error(f"Failed to process {file_path.name}: {e}")
        
        return documents
